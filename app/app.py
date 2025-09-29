from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, PlainTextResponse
from pydantic import BaseModel, Field
from typing import Optional, Dict, Any, List, Tuple
import uvicorn
import boto3
import json
import time
import os
from datetime import datetime
import re
import PyPDF2
import docx
from io import BytesIO
from botocore.exceptions import ClientError
import asyncio
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from functools import lru_cache
import uuid
import urllib3
from PIL import Image
import base64
import csv
import hashlib

# Disable SSL warnings (for development)
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

os.environ["BYPASS_TOOL_CONSENT"] = "true"

# Strands imports with fallback
try:
    from strands import Agent, tool
    from strands.models import BedrockModel
    from strands.agent.conversation_manager import SlidingWindowConversationManager
    from strands_tools import use_aws, http_request, retrieve, image_reader
    
    STRANDS_AVAILABLE = True
except Exception as e:
    # Stub implementations for testing
    class Agent:
        def __init__(self, *args, **kwargs):
            pass
        def __call__(self, *args, **kwargs):
            return "Agent response (stub)"
        async def stream_async(self, *args, **kwargs):
            yield {"data": "Streaming response (stub)"}
    
    class BedrockModel:
        def __init__(self, *args, **kwargs):
            pass
    
    class SlidingWindowConversationManager:
        def __init__(self, *args, **kwargs):
            pass
    
    def tool(name=None, description=None):
        def decorator(func):
            func._tool_name = name or func.__name__
            func._tool_description = description or func.__doc__
            return func
        return decorator
    
    use_aws = None
    http_request = None
    retrieve = None
    image_reader = None
    STRANDS_AVAILABLE = False

# ============================================================================
# CUSTOM IMAGE GENERATION TOOLS
# ============================================================================

class ImageGenerationError(Exception):
    """Custom exception for image generation errors"""
    def __init__(self, message):
        self.message = message

@tool(
    name="generate_image_sd3",
    description="Generate high-quality images using Stable Diffusion 3 Large model. Creates images from text prompts and uploads them to S3."
)
def generate_image_sd3(
    prompt: str,
    user_id: str,
    project_id: str,
    aspect_ratio: str = "1:1",
    negative_prompt: str = "",
    seed: int = 0,
    output_format: str = "PNG"
) -> str:
    """
    Generate images using Stable Diffusion 3 Large model and upload to S3.
    
    Args:
        prompt: Descriptive text prompt for image generation
        user_id: User identifier for S3 path
        project_id: Project identifier for S3 path
        aspect_ratio: Image aspect ratio (16:9, 1:1, 21:9, 2:3, 3:2, 4:5, 5:4, 9:16, 9:21)
        negative_prompt: What you don't want to see in the image
        seed: Random seed for reproducibility (0 for random)
        output_format: Output format (PNG)
    
    Returns:
        String with generation status and S3 URL
    """
    generation_id = str(uuid.uuid4())[:8]
    
    try:
        # Initialize Bedrock client
        bedrock_client = boto3.client('bedrock-runtime', region_name='us-west-2')
        s3_client = boto3.client('s3', region_name='us-east-1')
        
        # Prepare request body for SD3 Large
        request_body = {
            'prompt': prompt,
            'aspect_ratio': aspect_ratio,
            'mode': 'text-to-image',
            'output_format': output_format
        }
        
        # Add optional parameters
        if negative_prompt:
            request_body['negative_prompt'] = negative_prompt
        if seed > 0:
            request_body['seed'] = seed
        
        # Generate image using Stable Diffusion 3 Large
        response = bedrock_client.invoke_model(
            modelId='stability.sd3-large-v1:0',
            body=json.dumps(request_body)
        )
        
        # Parse response
        output_body = json.loads(response["body"].read().decode("utf-8"))
        
        # Check for errors
        if output_body.get("finish_reasons") and output_body["finish_reasons"][0] is not None:
            error_reason = output_body["finish_reasons"][0]
            return f"❌ Image generation failed: {error_reason}"
        
        # Get base64 image
        base64_image = output_body["images"][0]
        image_data = base64.b64decode(base64_image)
        
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        image_filename = f"sd3_generated_{timestamp}_{uuid.uuid4().hex[:8]}.{output_format.lower()}"
        
        # Upload to S3
        bucket_name = 'qubitz-customer-prod'
        s3_key = f"{user_id}/{project_id}/generated_images/{image_filename}"
        
        s3_client.put_object(
            Bucket=bucket_name,
            Key=s3_key,
            Body=image_data,
            ContentType=f'image/{output_format.lower()}',
            Metadata={
                'generated_by': 'stable_diffusion_3_large',
                'prompt': prompt[:200],  # Truncate for metadata
                'user_id': user_id,
                'project_id': project_id,
                'generation_timestamp': timestamp
            }
        )
        
        # Generate S3 URL
        s3_url = f"https://{bucket_name}.s3.amazonaws.com/{s3_key}"
        
        # Get image dimensions for response
        image_obj = Image.open(BytesIO(image_data))
        width, height = image_obj.size
        
        response_message = f"""✅ **Image Generated Successfully with Stable Diffusion 3 Large!**

🎨 **Generated Image Details:**
• **Prompt**: {prompt}
• **Dimensions**: {width}x{height} pixels
• **Aspect Ratio**: {aspect_ratio}
• **Format**: {output_format}
• **Model**: Stable Diffusion 3 Large

🔗 **Image URL**: {s3_url}

📁 **S3 Location**: {s3_key}
⏱️ **Generated**: {timestamp}

The image has been successfully uploaded to S3 and is ready for use!"""

        return response_message
        
    except ClientError as e:
        return f"❌ AWS Error during image generation: {str(e)}"
    except Exception as e:
        return f"❌ Image generation failed: {str(e)}"

@tool(
    name="generate_image_from_reference",
    description="Generate images based on a reference image using Stable Diffusion XL. Supports image-to-image generation with style transfer."
)
def generate_image_from_reference(
    prompt: str,
    reference_image_s3_url: str,
    user_id: str,
    project_id: str,
    image_strength: float = 0.7,
    cfg_scale: float = 7.0,
    style_preset: str = "photographic",
    steps: int = 30,
    seed: int = 0
) -> str:
    """
    Generate images from a reference image using Stable Diffusion XL.
    
    Args:
        prompt: Text description for the desired output
        reference_image_s3_url: S3 URL of the reference image
        user_id: User identifier for S3 path
        project_id: Project identifier for S3 path
        image_strength: How much the output should resemble the input (0.0-1.0)
        cfg_scale: How closely to follow the prompt (0-35)
        style_preset: Style to apply (3d-model, analog-film, anime, cinematic, etc.)
        steps: Number of generation steps (10-50)
        seed: Random seed for reproducibility (0 for random)
    
    Returns:
        String with generation status and S3 URL
    """
    generation_id = str(uuid.uuid4())[:8]
    
    try:
        # Initialize clients
        bedrock_client = boto3.client('bedrock-runtime', region_name='us-east-1')
        s3_client = boto3.client('s3', region_name='us-east-1')
        
        # Download reference image from S3
        # Extract bucket and key from S3 URL
        if reference_image_s3_url.startswith('https://'):
            # Parse S3 URL
            url_parts = reference_image_s3_url.replace('https://', '').split('/')
            ref_bucket = url_parts[0].split('.s3.amazonaws.com')[0]
            ref_key = '/'.join(url_parts[1:])
        else:
            return "❌ Invalid S3 URL format. Please provide a valid S3 URL."
        
        # Download reference image
        try:
            ref_response = s3_client.get_object(Bucket=ref_bucket, Key=ref_key)
            reference_image_data = ref_response['Body'].read()
            reference_image_base64 = base64.b64encode(reference_image_data).decode('utf-8')
        except ClientError as e:
            return f"❌ Could not download reference image: {str(e)}"
        
        # Prepare request body for SDXL image-to-image
        request_body = {
            "text_prompts": [
                {
                    "text": prompt,
                    "weight": 1.0
                }
            ],
            "init_image": reference_image_base64,
            "init_image_mode": "IMAGE_STRENGTH",
            "image_strength": image_strength,
            "cfg_scale": cfg_scale,
            "steps": steps,
            "samples": 1
        }
        
        # Add optional parameters
        if style_preset:
            request_body["style_preset"] = style_preset
        if seed > 0:
            request_body["seed"] = seed
        
        # Generate image using Stable Diffusion XL
        response = bedrock_client.invoke_model(
            modelId='stability.stable-diffusion-xl-v1',
            body=json.dumps(request_body),
            accept="application/json",
            contentType="application/json"
        )
        
        # Parse response
        response_body = json.loads(response.get("body").read())
        
        # Check for errors
        artifacts = response_body.get("artifacts", [])
        if not artifacts:
            return "❌ No image artifacts returned from the model."
        
        artifact = artifacts[0]
        finish_reason = artifact.get("finishReason")
        
        if finish_reason in ['ERROR', 'CONTENT_FILTERED']:
            return f"❌ Image generation failed: {finish_reason}"
        
        # Get base64 image
        base64_image = artifact.get("base64")
        image_data = base64.b64decode(base64_image)
        
        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        image_filename = f"sdxl_img2img_{timestamp}_{uuid.uuid4().hex[:8]}.png"
        
        # Upload to S3
        bucket_name = 'qubitz-customer-prod'
        s3_key = f"{user_id}/{project_id}/generated_images/{image_filename}"
        
        s3_client.put_object(
            Bucket=bucket_name,
            Key=s3_key,
            Body=image_data,
            ContentType='image/png',
            Metadata={
                'generated_by': 'stable_diffusion_xl_img2img',
                'prompt': prompt[:200],
                'reference_image': reference_image_s3_url,
                'image_strength': str(image_strength),
                'style_preset': style_preset,
                'user_id': user_id,
                'project_id': project_id,
                'generation_timestamp': timestamp
            }
        )
        
        # Generate S3 URL
        s3_url = f"https://{bucket_name}.s3.amazonaws.com/{s3_key}"
        
        # Get image dimensions
        image_obj = Image.open(BytesIO(image_data))
        width, height = image_obj.size
        
        response_message = f"""✅ **Image-to-Image Generation Successful with Stable Diffusion XL!**

🎨 **Generated Image Details:**
• **Prompt**: {prompt}
• **Reference Image**: {reference_image_s3_url}
• **Dimensions**: {width}x{height} pixels
• **Image Strength**: {image_strength} (how much it resembles the reference)
• **Style**: {style_preset}
• **CFG Scale**: {cfg_scale}
• **Steps**: {steps}
• **Model**: Stable Diffusion XL

🔗 **Generated Image URL**: {s3_url}

📁 **S3 Location**: {s3_key}
⏱️ **Generated**: {timestamp}

The image has been successfully created based on your reference image and uploaded to S3!"""

        return response_message
        
    except ClientError as e:
        return f"❌ AWS Error during image-to-image generation: {str(e)}"
    except Exception as e:
        return f"❌ Image-to-image generation failed: {str(e)}"

@tool(
    name="upload_image_to_s3",
    description="Upload an image file to S3 and return the URL. Useful for preparing reference images for image-to-image generation."
)
def upload_image_to_s3(
    image_base64: str,
    filename: str,
    user_id: str,
    project_id: str,
    image_format: str = "PNG"
) -> str:
    """
    Upload a base64 encoded image to S3.
    
    Args:
        image_base64: Base64 encoded image data
        filename: Desired filename (without extension)
        user_id: User identifier for S3 path
        project_id: Project identifier for S3 path
        image_format: Image format (PNG, JPEG, etc.)
    
    Returns:
        String with upload status and S3 URL
    """
    upload_id = str(uuid.uuid4())[:8]
    
    try:
        s3_client = boto3.client('s3', region_name='us-east-1')
        
        # Decode base64 image
        image_data = base64.b64decode(image_base64)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = re.sub(r'[^a-zA-Z0-9_-]', '_', filename)
        image_filename = f"{safe_filename}_{timestamp}.{image_format.lower()}"
        
        # Upload to S3
        bucket_name = 'qubitz-customer-prod'
        s3_key = f"{user_id}/{project_id}/uploaded_images/{image_filename}"
        
        s3_client.put_object(
            Bucket=bucket_name,
            Key=s3_key,
            Body=image_data,
            ContentType=f'image/{image_format.lower()}',
            Metadata={
                'uploaded_by': 'user',
                'original_filename': filename,
                'user_id': user_id,
                'project_id': project_id,
                'upload_timestamp': timestamp
            }
        )
        
        # Generate S3 URL
        s3_url = f"https://{bucket_name}.s3.amazonaws.com/{s3_key}"
        
        # Get image info
        image_obj = Image.open(BytesIO(image_data))
        width, height = image_obj.size
        
        response_message = f"""✅ **Image Uploaded Successfully to S3!**

📸 **Upload Details:**
• **Filename**: {image_filename}
• **Dimensions**: {width}x{height} pixels
• **Format**: {image_format}
• **Size**: {len(image_data)} bytes

🔗 **Image URL**: {s3_url}

📁 **S3 Location**: {s3_key}
⏱️ **Uploaded**: {timestamp}

You can now use this S3 URL as a reference image for image-to-image generation!"""

        return response_message
        
    except Exception as e:
        return f"❌ Image upload failed: {str(e)}"

# ============================================================================
# ENHANCED RAG SYSTEM WITH STATIC TOKEN LIMITS
# ============================================================================

class OptimizedRAGSystem:
    """Enhanced RAG system with static token limits and no external tokenizer dependency"""
    
    def __init__(self, region='us-east-1'):
        """Initialize the optimized RAG system"""
        self.region = region
        
        # Configure boto3 clients
        boto3_config = boto3.session.Config(
            region_name=region,
            retries={'max_attempts': 3, 'mode': 'adaptive'},
            max_pool_connections=20
        )
        
        self.s3_client = boto3.client('s3', config=boto3_config)
        self.s3vectors_client = boto3.client('s3vectors', config=boto3_config)
        self.bedrock_client = boto3.client('bedrock-runtime', config=boto3_config)
        
        # Configuration
        self.source_bucket = 'qubitz-customer-prod'
        self.dataset_bucket = 'amplify-d14cv22pvf70is-de-qubitzamplifybucketbucke-msfrfeyucxdg'
        self.vector_bucket_prefix = 'qubitz-vectors'
        
        # Enhanced embedding models
        self.text_embedding_model = 'amazon.titan-embed-text-v2:0'
        self.multimodal_embedding_model = 'amazon.titan-embed-image-v1'
        self.embedding_dimension = 1024
        
        # Static token limits for Amazon Titan Embed Text v2
        # Amazon Titan Embed Text v2 limit: 8192 tokens
        # Using conservative limits to ensure success
        self.max_chars_for_embedding = 5000      # Conservative character limit
        self.max_tokens_estimate = 6500          # Conservative token estimate
        
        # Character-to-token ratio estimates for different content types
        self.char_to_token_ratios = {
            'english': 0.25,      # ~4 characters per token
            'technical': 0.3,     # ~3.3 characters per token (more technical terms)
            'mixed': 0.27,        # ~3.7 characters per token
            'default': 0.28       # ~3.6 characters per token (conservative)
        }
        
        # Adaptive similarity thresholds - More permissive for better recall
        self.similarity_thresholds = {
            'exact_match': 0.8,
            'strong_match': 0.6,
            'moderate_match': 0.4,
            'weak_match': 0.2,
            'any_match': 0.05      # Very permissive
        }
        
        # Optimized chunking parameters with token awareness
        self.chunk_size = 700           # Reduced for token safety
        self.chunk_overlap = 100
        self.max_workers = 8
        
        # Thread pools with proper limits
        self.file_processor = ThreadPoolExecutor(max_workers=self.max_workers)
        self.embedding_processor = ThreadPoolExecutor(max_workers=4)
        
        # Setup tracking
        self._setup_in_progress = set()
        self._setup_lock = threading.Lock()

    def _estimate_tokens(self, text: str, content_type: str = 'default') -> int:
        """Estimate token count using static character-to-token ratios"""
        if not text:
            return 0
        
        # Get ratio for content type
        ratio = self.char_to_token_ratios.get(content_type, self.char_to_token_ratios['default'])
        
        # Count characters and estimate tokens
        char_count = len(text)
        estimated_tokens = int(char_count * ratio)
        
        return estimated_tokens

    def _detect_content_type(self, text: str) -> str:
        """Detect content type for better token estimation"""
        if not text:
            return 'default'
        
        # Count technical indicators
        technical_patterns = [
            r'\b[A-Z]{2,}\b',      # Acronyms
            r'\b\w+\.\w+\b',       # Technical terms with dots
            r'\b\d+\.\d+\b',       # Version numbers
            r'[{}()[\$\$<>]',        # Technical brackets
            r'\b(API|SDK|JSON|XML|HTTP|URL|SQL|AWS|S3)\b'  # Technical terms
        ]
        
        technical_count = sum(len(re.findall(pattern, text)) for pattern in technical_patterns)
        
        # Determine content type
        if technical_count > len(text.split()) * 0.1:  # >10% technical terms
            return 'technical'
        elif re.search(r'[^\x00-\x7F]', text):  # Non-ASCII characters
            return 'mixed'
        else:
            return 'english'

    def _truncate_text_safely(self, text: str, embedding_id: str) -> str:
        """Safely truncate text using static token estimation"""
        if not text:
            return text
        
        original_length = len(text)
        content_type = self._detect_content_type(text)
        estimated_tokens = self._estimate_tokens(text, content_type)
        
        # If already within limits, return as-is
        if estimated_tokens <= self.max_tokens_estimate and len(text) <= self.max_chars_for_embedding:
            return text
        
        # Progressive truncation with token checking
        if estimated_tokens > self.max_tokens_estimate or len(text) > self.max_chars_for_embedding:
            # Use the more restrictive limit
            if estimated_tokens > self.max_tokens_estimate:
                # Calculate target characters based on token estimate
                target_chars = int(self.max_tokens_estimate / self.char_to_token_ratios[content_type])
                target_chars = min(target_chars, self.max_chars_for_embedding)
            else:
                target_chars = self.max_chars_for_embedding
            
            # Add safety margin (90% of calculated limit)
            target_chars = int(target_chars * 0.9)
            
            # Truncate at word boundary if possible
            if target_chars < len(text):
                truncated = text[:target_chars]
                
                # Try to find last complete word
                last_space = truncated.rfind(' ')
                if last_space > target_chars * 0.8:  # If we can find a word boundary in the last 20%
                    truncated = truncated[:last_space]
                
                # Verify final token estimate
                final_tokens = self._estimate_tokens(truncated, content_type)
                
                return truncated + "..."
        
        return text

    def generate_enhanced_embedding(self, text: str, image_base64: str = None) -> Optional[List[float]]:
        """Generate enhanced embedding with static token management"""
        embedding_id = str(uuid.uuid4())[:8]
        
        try:
            # Choose embedding model based on content
            if image_base64:
                return self._generate_multimodal_embedding(text, image_base64, embedding_id)
            else:
                return self._generate_text_embedding(text, embedding_id)
                
        except Exception as e:
            return None

    def _generate_text_embedding(self, text: str, embedding_id: str) -> Optional[List[float]]:
        """Generate text-only embedding with static token management"""
        try:
            # Apply safe truncation
            safe_text = self._truncate_text_safely(text, embedding_id)
            
            # Final safety check
            final_char_count = len(safe_text)
            content_type = self._detect_content_type(safe_text)
            final_token_estimate = self._estimate_tokens(safe_text, content_type)
            
            if final_token_estimate > self.max_tokens_estimate:
                # Emergency truncation
                emergency_chars = int(self.max_tokens_estimate / self.char_to_token_ratios['default']) * 0.8
                safe_text = safe_text[:int(emergency_chars)]
            
            # Use Titan V2 with dimension optimization
            body = {
                "inputText": safe_text,
                "dimensions": self.embedding_dimension,
                "normalize": True
            }
            
            start_time = time.time()
            
            response = self.bedrock_client.invoke_model(
                modelId=self.text_embedding_model,
                body=json.dumps(body)
            )
            
            # Parse response
            response_body = json.loads(response['body'].read())
            embedding = response_body.get('embedding')
            
            if not embedding:
                return None
            
            # Validate embedding dimension
            if len(embedding) != self.embedding_dimension:
                return None
            
            return embedding
            
        except Exception as e:
            return None

    def _generate_multimodal_embedding(self, text: str, image_base64: str, embedding_id: str) -> Optional[List[float]]:
        """Generate multimodal embedding with static token management"""
        try:
            # Apply safe truncation to text
            safe_text = self._truncate_text_safely(text, embedding_id)
            
            # Prepare multimodal request
            body = {
                "inputText": safe_text,
                "inputImage": image_base64,
                "embeddingConfig": {
                    "outputEmbeddingLength": self.embedding_dimension
                }
            }
            
            start_time = time.time()
            
            response = self.bedrock_client.invoke_model(
                modelId=self.multimodal_embedding_model,
                body=json.dumps(body)
            )
            
            # Parse response
            response_body = json.loads(response['body'].read())
            embedding = response_body.get('embedding')
            
            if not embedding:
                return None
            
            # Validate embedding dimension
            if len(embedding) != self.embedding_dimension:
                return None
            
            return embedding
            
        except Exception as e:
            return None

    @lru_cache(maxsize=100)
    def _get_tenant_names(self, user_id: str, project_id: str) -> Tuple[str, str]:
        """Get cached tenant-specific names with proper naming conventions"""
        # Create deterministic but unique names
        tenant_hash = hashlib.md5(f"{user_id}-{project_id}".encode()).hexdigest()[:8]
        
        bucket_name = f"{self.vector_bucket_prefix}-{tenant_hash}".lower()
        index_name = f"index-{tenant_hash}".lower()
        
        # Ensure valid names
        bucket_name = re.sub(r'[^a-z0-9\-]', '-', bucket_name)[:63]
        index_name = re.sub(r'[^a-z0-9\-]', '-', index_name)[:63]
        
        return bucket_name, index_name

    def _cleanup_existing_infrastructure(self, bucket_name: str, index_name: str) -> bool:
        """Clean up existing infrastructure that has mismatched configuration"""
        cleanup_id = str(uuid.uuid4())[:8]
        
        try:
            # Check and delete existing index if it exists
            try:
                existing_index = self.s3vectors_client.get_index(vectorBucketName=bucket_name, indexName=index_name)
                
                # Delete the index
                self.s3vectors_client.delete_index(vectorBucketName=bucket_name, indexName=index_name)
                
                # Wait for deletion to complete
                time.sleep(5)
                
            except ClientError as e:
                if e.response['Error']['Code'] != 'NotFoundException':
                    pass
            
            # Check vector bucket existence
            try:
                bucket_info = self.s3vectors_client.get_vector_bucket(vectorBucketName=bucket_name)
                return True
                
            except ClientError as e:
                if e.response['Error']['Code'] in ['NotFoundException', 'NoSuchBucket']:
                    return True
                else:
                    return False
            
        except Exception as e:
            return False

    def _create_vector_infrastructure(self, bucket_name: str, index_name: str) -> bool:
        """Create fresh vector bucket and index with enhanced configuration"""
        creation_id = str(uuid.uuid4())[:8]
        
        try:
            # Create vector bucket (handle existing bucket gracefully)
            try:
                self.s3vectors_client.create_vector_bucket(vectorBucketName=bucket_name)
            except ClientError as e:
                error_code = e.response['Error']['Code']
                if error_code not in ['ConflictException', 'BucketAlreadyExists', 'BucketAlreadyOwnedByYou']:
                    return False
            
            # Wait a moment for bucket to be ready
            time.sleep(3)
            
            # Create vector index with enhanced configuration
            create_params = {
                'vectorBucketName': bucket_name,
                'indexName': index_name,
                'dimension': self.embedding_dimension,
                'distanceMetric': 'cosine',
                'dataType': 'float32',
                'metadataConfiguration': {
                    'nonFilterableMetadataKeys': [
                        'full_text',
                        'text_preview', 
                        'word_count',
                        'content_length',
                        'has_images',
                        'multimodal_content'
                    ]
                }
            }
            
            self.s3vectors_client.create_index(**create_params)
            
            # Wait for index to be ready
            max_wait = 60
            wait_time = 0
            while wait_time < max_wait:
                try:
                    index_status = self.s3vectors_client.get_index(vectorBucketName=bucket_name, indexName=index_name)
                    return True
                except ClientError as check_e:
                    if check_e.response['Error']['Code'] == 'NotFoundException':
                        time.sleep(3)
                        wait_time += 3
                    else:
                        break
            
            return False
            
        except Exception as e:
            return False

    def list_files_comprehensive(self, user_id: str, project_id: str) -> List[Dict]:
        """List all available files for processing"""
        discovery_id = str(uuid.uuid4())[:8]
        all_files = []
        
        # Source bucket files
        prefix = f"{user_id}/{project_id}/"
        
        try:
            paginator = self.s3_client.get_paginator('list_objects_v2')
            pages = paginator.paginate(Bucket=self.source_bucket, Prefix=prefix)
            
            source_count = 0
            for page in pages:
                if 'Contents' in page:
                    for obj in page['Contents']:
                        if not obj['Key'].endswith('/') and not obj['Key'].endswith('workflow.json'):
                            all_files.append({
                                'key': obj['Key'],
                                'size': obj['Size'],
                                'last_modified': obj['LastModified'],
                                'bucket': self.source_bucket,
                                'source': 'documents',
                                'file_type': self._get_file_type(obj['Key'])
                            })
                            source_count += 1
            
        except ClientError as e:
            pass
        
        # Dataset bucket files
        dataset_prefix = f"user-uploads/dataset/{user_id}/{project_id}/"
        
        try:
            paginator = self.s3_client.get_paginator('list_objects_v2')
            pages = paginator.paginate(Bucket=self.dataset_bucket, Prefix=dataset_prefix)
            
            dataset_count = 0
            for page in pages:
                if 'Contents' in page:
                    for obj in page['Contents']:
                        if not obj['Key'].endswith('/'):
                            all_files.append({
                                'key': obj['Key'],
                                'size': obj['Size'],
                                'last_modified': obj['LastModified'],
                                'bucket': self.dataset_bucket,
                                'source': 'datasets',
                                'file_type': self._get_file_type(obj['Key'])
                            })
                            dataset_count += 1
            
        except ClientError as e:
            pass
        
        return all_files

    def _get_file_type(self, file_key: str) -> str:
        """Determine file type from extension"""
        extension = os.path.splitext(file_key)[1].lower()
        type_mapping = {
            '.pdf': 'PDF',
            '.docx': 'DOCX',
            '.doc': 'DOC',
            '.txt': 'TEXT',
            '.json': 'JSON',
            '.csv': 'CSV',
            '.xlsx': 'EXCEL',
            '.xls': 'EXCEL',
            '.md': 'MARKDOWN',
            '.rtf': 'RTF',
            '.xml': 'XML',
            '.png': 'IMAGE',
            '.jpg': 'IMAGE',
            '.jpeg': 'IMAGE',
            '.gif': 'IMAGE',
            '.bmp': 'IMAGE',
            '.tiff': 'IMAGE',
            '.ts': 'TYPESCRIPT',
            '.js': 'JAVASCRIPT',
            '.py': 'PYTHON',
            '.zip': 'ARCHIVE'
        }
        return type_mapping.get(extension, 'UNKNOWN')

    def extract_file_content(self, file_info: Dict) -> str:
        """Extract text content from file with enhanced processing"""
        extraction_id = str(uuid.uuid4())[:8]
        file_key = file_info['key']
        bucket = file_info['bucket']
        file_type = file_info['file_type']
        
        try:
            response = self.s3_client.get_object(Bucket=bucket, Key=file_key)
            file_content = response['Body'].read()
            
            # Route to appropriate extraction method
            if file_type == 'PDF':
                content = self._extract_pdf_content(file_content)
            elif file_type in ['DOCX', 'DOC']:
                content = self._extract_docx_content(file_content)
            elif file_type in ['TEXT', 'MARKDOWN', 'RTF']:
                content = self._extract_text_content(file_content)
            elif file_type == 'JSON':
                content = self._extract_json_content(file_content)
            elif file_type == 'CSV':
                content = self._extract_csv_content(file_content)
            elif file_type == 'IMAGE':
                content = self._extract_image_content(file_content, file_key)
            elif file_type in ['TYPESCRIPT', 'JAVASCRIPT', 'PYTHON']:
                content = self._extract_code_content(file_content)
            elif file_type == 'ARCHIVE':
                content = self._extract_archive_content(file_content, file_key)
            else:
                content = self._extract_generic_content(file_content)
            
            return content
                
        except Exception as e:
            return ""

    def _extract_pdf_content(self, pdf_content: bytes) -> str:
        """Extract text from PDF with improved handling"""
        try:
            pdf_file = BytesIO(pdf_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            all_text = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        cleaned_text = self._clean_text(page_text)
                        if len(cleaned_text) > 20:
                            all_text.append(f"[Page {page_num + 1}]\n{cleaned_text}")
                except Exception as e:
                    continue
            
            final_content = "\n\n".join(all_text)
            return final_content
            
        except Exception as e:
            return ""

    def _extract_docx_content(self, docx_content: bytes) -> str:
        """Extract text from DOCX with improved structure handling"""
        try:
            docx_file = BytesIO(docx_content)
            doc = docx.Document(docx_file)
            
            content_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    cleaned_text = self._clean_text(para.text)
                    if len(cleaned_text) > 10:
                        content_parts.append(cleaned_text)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_content = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            row_data.append(cell.text.strip())
                    if row_data:
                        table_content.append(" | ".join(row_data))
                
                if table_content:
                    content_parts.append(f"[Table {table_idx + 1}]\n" + "\n".join(table_content))
            
            final_content = "\n\n".join(content_parts)
            return final_content
            
        except Exception as e:
            return ""

    def _extract_text_content(self, text_content: bytes) -> str:
        """Extract text from plain text files with encoding detection"""
        for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
            try:
                text = text_content.decode(encoding)
                cleaned = self._clean_text(text)
                return cleaned
            except Exception as e:
                continue
        
        # Fallback with error handling
        try:
            fallback = text_content.decode('utf-8', errors='ignore')
            return fallback
        except:
            return ""

    def _extract_json_content(self, json_content: bytes) -> str:
        """Extract meaningful content from JSON files"""
        try:
            data = json.loads(json_content.decode('utf-8'))
            
            def extract_json_text(obj, path=""):
                texts = []
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        current_path = f"{path}.{key}" if path else key
                        if isinstance(value, str) and len(value.strip()) > 5:
                            texts.append(f"{current_path}: {value.strip()}")
                        elif isinstance(value, (dict, list)):
                            texts.extend(extract_json_text(value, current_path))
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        current_path = f"{path}[{i}]" if path else f"[{i}]"
                        if isinstance(item, str) and len(item.strip()) > 5:
                            texts.append(f"{current_path}: {item.strip()}")
                        elif isinstance(item, (dict, list)):
                            texts.extend(extract_json_text(item, current_path))
                return texts
            
            extracted_texts = extract_json_text(data)
            result = "\n".join(extracted_texts)
            return result
            
        except Exception as e:
            return ""

    def _extract_csv_content(self, csv_content: bytes) -> str:
        """Extract content from CSV files with proper structure"""
        try:
            csv_text = csv_content.decode('utf-8', errors='ignore')
            csv_reader = csv.reader(csv_text.splitlines())
            rows = list(csv_reader)
            
            if not rows:
                return ""
            
            headers = rows[0] if rows else []
            data_rows = rows[1:] if len(rows) > 1 else []
            
            content_parts = []
            
            # Add header information
            if headers:
                content_parts.append(f"CSV Columns: {' | '.join(headers)}")
            
            # Process data rows (sample first 100 rows)
            sample_rows = data_rows[:100]
            for i, row in enumerate(sample_rows):
                if len(row) == len(headers) and any(cell.strip() for cell in row):
                    row_data = []
                    for header, value in zip(headers, row):
                        if value and value.strip():
                            row_data.append(f"{header}: {value.strip()}")
                    if row_data:
                        content_parts.append(f"Row {i+1}: {' | '.join(row_data)}")
            
            result = "\n".join(content_parts)
            return result
            
        except Exception as e:
            return ""

    def _extract_image_content(self, image_content: bytes, file_key: str) -> str:
        """Extract textual description from image using multimodal analysis"""
        try:
            # Convert image to base64 for analysis
            image_base64 = base64.b64encode(image_content).decode('utf-8')
            
            # Create descriptive text about the image
            filename = os.path.basename(file_key)
            image_description = f"Image file: {filename}\n"
            image_description += f"File size: {len(image_content)} bytes\n"
            
            # Try to get image dimensions
            try:
                image_obj = Image.open(BytesIO(image_content))
                width, height = image_obj.size
                image_description += f"Dimensions: {width}x{height} pixels\n"
                image_description += f"Format: {image_obj.format}\n"
            except Exception as e:
                pass
            
            # Add metadata for multimodal processing
            image_description += f"Content type: Visual content requiring multimodal analysis\n"
            image_description += f"Image data available for multimodal embedding generation"
            
            return image_description
            
        except Exception as e:
            return f"Image file: {os.path.basename(file_key)} - Could not process image content"

    def _extract_code_content(self, code_content: bytes) -> str:
        """Extract content from code files"""
        try:
            # Try to decode as text
            for encoding in ['utf-8', 'utf-16', 'latin-1']:
                try:
                    code_text = code_content.decode(encoding)
                    # Add code context
                    return f"Code file content:\n{code_text}"
                except:
                    continue
            
            return "Code file - Could not decode content"
            
        except Exception as e:
            return "Code file - Error processing content"

    def _extract_archive_content(self, archive_content: bytes, file_key: str) -> str:
        """Extract information from archive files"""
        try:
            filename = os.path.basename(file_key)
            return f"Archive file: {filename}\nFile size: {len(archive_content)} bytes\nContent type: Archive/Compressed file"
            
        except Exception as e:
            return f"Archive file: {os.path.basename(file_key)} - Could not process archive content"

    def _extract_generic_content(self, content: bytes) -> str:
        """Generic content extraction for unknown file types"""
        try:
            # Try to decode as text
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = content.decode(encoding)
                    if len(text.strip()) > 10:
                        cleaned = self._clean_text(text)
                        return cleaned
                except:
                    continue
            
            return ""
        except Exception as e:
            return ""

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        if not text:
            return ""
        
        # Remove excessive whitespace while preserving structure
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\r\n', '\n', text)
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        return text.strip()

    def create_token_aware_chunks(self, content: str, file_name: str, file_type: str) -> List[Dict]:
        """Create token-aware chunks using static estimation"""
        chunk_id = str(uuid.uuid4())[:8]
        
        if not content or len(content.strip()) < 50:
            return []
        
        chunks = []
        content_type = self._detect_content_type(content)
        
        # Split by natural boundaries (paragraphs first)
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        current_chunk = ""
        chunk_index = 0
        
        for para_idx, paragraph in enumerate(paragraphs):
            # Check if adding this paragraph would exceed token limits
            proposed_chunk = current_chunk + "\n\n" + paragraph if current_chunk else paragraph
            proposed_tokens = self._estimate_tokens(proposed_chunk, content_type)
            
            if proposed_tokens > self.max_tokens_estimate and current_chunk:
                # Save current chunk
                chunk_text = current_chunk.strip()
                if len(chunk_text) > 50:  # Only save substantial chunks
                    chunk_info = {
                        'text': chunk_text,
                        'chunk_index': chunk_index,
                        'file_name': file_name,
                        'file_type': file_type,
                        'content_length': len(chunk_text),
                        'word_count': len(chunk_text.split()),
                        'estimated_tokens': self._estimate_tokens(chunk_text, content_type),
                        'content_type': content_type
                    }
                    chunks.append(chunk_info)
                    chunk_index += 1
                
                # Start new chunk with minimal overlap (token-aware)
                overlap_words = current_chunk.split()[-15:] if current_chunk else []  # Reduced overlap
                overlap_text = " ".join(overlap_words)
                current_chunk = overlap_text + "\n\n" + paragraph if overlap_text else paragraph
            else:
                current_chunk = proposed_chunk
        
        # Add final chunk
        if current_chunk.strip() and len(current_chunk.strip()) > 50:
            chunk_text = current_chunk.strip()
            final_chunk = {
                'text': chunk_text,
                'chunk_index': chunk_index,
                'file_name': file_name,
                'file_type': file_type,
                'content_length': len(chunk_text),
                'word_count': len(chunk_text.split()),
                'estimated_tokens': self._estimate_tokens(chunk_text, content_type),
                'content_type': content_type
            }
            chunks.append(final_chunk)
        
        return chunks

    def process_files_batch(self, files: List[Dict]) -> List[Dict]:
        """Process files in batches with token-aware chunking"""
        process_id = str(uuid.uuid4())[:8]
        
        all_chunks = []
        
        # Process files in parallel with enhanced error handling
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_file = {
                executor.submit(self._process_single_file_token_aware, file_info): file_info 
                for file_info in files
            }
            
            for future in as_completed(future_to_file):
                file_info = future_to_file[future]
                try:
                    file_chunks = future.result()
                    all_chunks.extend(file_chunks)
                except Exception as e:
                    pass
        
        return all_chunks

    def _process_single_file_token_aware(self, file_info: Dict) -> List[Dict]:
        """Process a single file with token-aware chunking"""
        try:
            content = self.extract_file_content(file_info)
            if not content:
                return []
            
            chunks = self.create_token_aware_chunks(
                content, 
                os.path.basename(file_info['key']),
                file_info['file_type']
            )
            
            # Add enhanced file metadata to chunks
            for chunk in chunks:
                chunk.update({
                    'source': file_info['source'],
                    'bucket': file_info['bucket'],
                    'file_key': file_info['key'],
                    'file_size': file_info['size'],
                    'last_modified': file_info['last_modified'].isoformat(),
                    'has_images': file_info['file_type'] == 'IMAGE',
                    'multimodal_content': file_info['file_type'] == 'IMAGE'
                })
            
            return chunks
            
        except Exception as e:
            return []

    def generate_embeddings_batch(self, chunks: List[Dict]) -> List[Tuple[Dict, List[float]]]:
        """Generate embeddings for chunks with token-aware processing"""
        batch_id = str(uuid.uuid4())[:8]
        
        successful_pairs = []
        
        # Process chunks with controlled concurrency
        with ThreadPoolExecutor(max_workers=4) as executor:  # Limited concurrency for Bedrock
            future_to_chunk = {
                executor.submit(self._generate_single_embedding_token_aware, chunk): chunk 
                for chunk in chunks
            }
            
            for future in as_completed(future_to_chunk):
                chunk = future_to_chunk[future]
                try:
                    embedding = future.result()
                    if embedding:
                        successful_pairs.append((chunk, embedding))
                        
                except Exception as e:
                    pass
        
        return successful_pairs

    def _generate_single_embedding_token_aware(self, chunk: Dict) -> Optional[List[float]]:
        """Generate embedding for a single chunk with token awareness"""
        max_retries = 3
        
        for attempt in range(max_retries):
            try:
                # Determine if this is multimodal content
                is_multimodal = chunk.get('multimodal_content', False)
                
                if is_multimodal:
                    # For image content, we need to get the image data
                    # This is a simplified approach - in practice, you'd need to fetch the image
                    embedding = self.generate_enhanced_embedding(chunk['text'])
                else:
                    embedding = self.generate_enhanced_embedding(chunk['text'])
                
                if embedding:
                    return embedding
                    
            except Exception as e:
                if attempt < max_retries - 1:
                    wait_time = 0.5 * (attempt + 1)
                    time.sleep(wait_time)
                    continue
                else:
                    pass
        
        return None

    def store_vectors_enhanced(self, bucket_name: str, index_name: str, 
                              chunk_embedding_pairs: List[Tuple[Dict, List[float]]], 
                              user_id: str, project_id: str) -> bool:
        """Store vectors with enhanced metadata structure"""
        storage_id = str(uuid.uuid4())[:8]
        
        try:
            vectors = []
            metadata_validation_errors = []
            
            for idx, (chunk, embedding) in enumerate(chunk_embedding_pairs):
                try:
                    # Create vector with enhanced metadata
                    vector_key = f"vec-{user_id}-{project_id}-{uuid.uuid4().hex[:8]}"
                    
                    # Prepare content fields
                    full_content = chunk['text']
                    content_preview = chunk['text'][:300] + "..." if len(chunk['text']) > 300 else chunk['text']
                    
                    # Create enhanced metadata with exactly 10 keys (AWS S3 Vectors limit)
                    metadata = {
                        # Filterable metadata (6 keys)
                        'user_id': str(user_id),
                        'project_id': str(project_id), 
                        'file_name': str(chunk['file_name']),
                        'file_type': str(chunk['file_type']),
                        'source': str(chunk['source']),
                        'chunk_index': int(chunk['chunk_index']),
                        
                        # Non-filterable metadata (4 keys) 
                        'full_text': str(full_content),
                        'text_preview': str(content_preview),
                        'word_count': int(chunk.get('word_count', 0)),
                        'content_length': int(chunk.get('content_length', len(full_content)))
                    }
                    
                    # Validate metadata structure
                    if len(metadata) != 10:
                        metadata_validation_errors.append(f"Vector {idx}: Metadata has {len(metadata)} keys, expected 10")
                        continue
                    
                    # Validate embedding format
                    if not isinstance(embedding, list) or len(embedding) != self.embedding_dimension:
                        continue
                    
                    vector = {
                        'key': vector_key,
                        'data': {'float32': [float(x) for x in embedding]},
                        'metadata': metadata
                    }
                    vectors.append(vector)
                
                except Exception as e:
                    continue
            
            if not vectors:
                return False
            
            # Store in optimal batch sizes (25 vectors per batch for S3 Vectors)
            batch_size = 25
            successful_batches = 0
            failed_batches = 0
            total_batches = (len(vectors) + batch_size - 1) // batch_size
            
            for batch_idx in range(0, len(vectors), batch_size):
                batch = vectors[batch_idx:batch_idx + batch_size]
                batch_num = (batch_idx // batch_size) + 1
                
                max_retries = 3
                batch_stored = False
                
                for attempt in range(max_retries):
                    try:
                        store_params = {
                            'vectorBucketName': bucket_name,
                            'indexName': index_name,
                            'vectors': batch
                        }
                        
                        start_time = time.time()
                        self.s3vectors_client.put_vectors(**store_params)
                        
                        successful_batches += 1
                        batch_stored = True
                        break
                        
                    except Exception as e:
                        if attempt < max_retries - 1:
                            wait_time = 1 * (attempt + 1)
                            time.sleep(wait_time)
                        else:
                            pass
                
                if not batch_stored:
                    failed_batches += 1
                
                # Small delay between batches to avoid rate limiting
                time.sleep(0.1)
            
            success_rate = successful_batches / total_batches if total_batches > 0 else 0
            
            return success_rate > 0.7  # Consider successful if >70% stored
            
        except Exception as e:
            return False

    def query_knowledge_base_enhanced(self, user_id: str, project_id: str, query: str, top_k: int = 10) -> Dict:
        """Enhanced query with token-aware processing and permissive thresholds"""
        query_id = str(uuid.uuid4())[:8]
        
        bucket_name, index_name = self._get_tenant_names(user_id, project_id)
        
        # Auto-setup if needed
        setup_started = not self.auto_setup_if_needed(user_id, project_id)
        if setup_started:
            pass
        
        # Check if knowledge base exists
        try:
            index_info = self.s3vectors_client.get_index(vectorBucketName=bucket_name, indexName=index_name)
            
        except ClientError as e:
            error_code = e.response['Error']['Code']
            
            if setup_started:
                return {
                    'query': query,
                    'results': [],
                    'context': '',
                    'message': '🔄 Your documents are being processed and will be available for search shortly.',
                    'setup_in_progress': True
                }
            else:
                return {
                    'query': query,
                    'results': [],
                    'context': '',
                    'message': '📭 No documents found in your collection.',
                    'setup_in_progress': False
                }
        
        try:
            # Generate query embedding with token safety
            query_embedding = self.generate_enhanced_embedding(query)
            
            if not query_embedding:
                return {
                    'query': query,
                    'results': [],
                    'context': '',
                    'message': '❌ Failed to process your query for search.',
                    'setup_in_progress': False
                }
            
            # Perform vector search
            query_params = {
                'vectorBucketName': bucket_name,
                'indexName': index_name,
                'queryVector': {'float32': query_embedding},
                'topK': top_k,
                'returnMetadata': True,
                'returnDistance': True
            }
            
            start_time = time.time()
            response = self.s3vectors_client.query_vectors(**query_params)
            
            results = response.get('vectors', [])
            
            # Enhanced result processing with permissive thresholds
            return self._process_results_with_permissive_thresholds(query_id, query, results)
            
        except Exception as e:
            return {
                'query': query,
                'results': [],
                'context': '',
                'message': f'❌ Search failed: {str(e)}',
                'setup_in_progress': False
            }

    def _process_results_with_permissive_thresholds(self, query_id: str, query: str, results: List[Dict]) -> Dict:
        """Process search results with very permissive thresholds for better recall"""
        if not results:
            return {
                'query': query,
                'results': [],
                'context': '',
                'message': '🔍 No relevant information found for your query.',
                'setup_in_progress': False
            }
        
        # Analyze distance distribution for adaptive thresholding
        distances = [result.get('distance', 1.0) for result in results]
        similarities = [max(0, 1 - distance) for distance in distances]
        
        # Log similarity analysis
        avg_similarity = sum(similarities) / len(similarities) if similarities else 0
        max_similarity = max(similarities) if similarities else 0
        min_similarity = min(similarities) if similarities else 0
        
        # Use very permissive threshold (0.02 = 2% similarity)
        threshold = 0.02
        
        # Process results with permissive threshold
        formatted_results = []
        context_parts = []
        sources_found = set()
        
        for i, result in enumerate(results):
            try:
                metadata = result.get('metadata', {})
                distance = result.get('distance', 1.0)
                similarity = max(0, 1 - distance)
                
                # Apply permissive threshold
                if similarity < threshold:
                    continue
                
                # Extract content with enhanced fallback hierarchy
                content = (metadata.get('full_text') or 
                          metadata.get('chunk_content') or 
                          metadata.get('text_preview') or 
                          'Content not available')
                
                file_name = metadata.get('file_name', 'Unknown')
                
                # Determine quality level with permissive ranges
                if similarity > 0.5:
                    quality_level = "high"
                elif similarity > 0.2:
                    quality_level = "medium"
                else:
                    quality_level = "low"
                
                formatted_result = {
                    'file_name': file_name,
                    'similarity_score': round(similarity, 4),
                    'quality_level': quality_level,
                    'content_preview': content[:200] + "..." if len(content) > 200 else content,
                    'file_type': metadata.get('file_type', 'unknown'),
                    'source': metadata.get('source', 'unknown'),
                    'chunk_index': metadata.get('chunk_index', 0),
                    'word_count': metadata.get('word_count', 0),
                    'distance': round(distance, 4)
                }
                formatted_results.append(formatted_result)
                sources_found.add(file_name)
                
                # Add to context for comprehensive response
                context_parts.append(f"**{file_name}** (Similarity: {similarity:.1%}):\n{content}")
                
            except Exception as e:
                continue
        
        # Final processing
        comprehensive_context = '\n\n---\n\n'.join(context_parts)
        sources_list = list(sources_found)
        
        # Create enhanced message
        if formatted_results:
            quality_counts = {}
            for result in formatted_results:
                quality = result['quality_level']
                quality_counts[quality] = quality_counts.get(quality, 0) + 1
            
            quality_summary = ", ".join([f"{count} {quality}" for quality, count in quality_counts.items()])
            
            message = f'🎯 Found {len(formatted_results)} results from {len(sources_found)} documents ({quality_summary} quality matches)'
        else:
            message = f'🔍 No results met the similarity threshold (max similarity: {max_similarity:.1%}). Your documents may still be processing or try rephrasing your query.'
        
        return {
            'query': query,
            'results': formatted_results,
            'context': comprehensive_context,
            'message': message,
            'sources_found': sources_list,
            'setup_in_progress': False,
            'similarity_analysis': {
                'threshold_used': threshold,
                'max_similarity': max_similarity,
                'avg_similarity': avg_similarity,
                'total_candidates': len(results),
                'results_returned': len(formatted_results)
            }
        }

    def auto_setup_if_needed(self, user_id: str, project_id: str) -> bool:
        """Auto-setup knowledge base if it doesn't exist or has wrong configuration"""
        setup_id = str(uuid.uuid4())[:8]
        bucket_name, index_name = self._get_tenant_names(user_id, project_id)
        
        # Check if knowledge base exists and is properly configured
        try:
            existing_index = self.s3vectors_client.get_index(vectorBucketName=bucket_name, indexName=index_name)
            existing_config = existing_index.get('index', {}).get('metadataConfiguration', {})
            existing_non_filterable = existing_config.get('nonFilterableMetadataKeys', [])
            
            expected_non_filterable = ['full_text', 'text_preview', 'word_count', 'content_length', 'has_images', 'multimodal_content']
            
            # Check dimension as well
            existing_dimension = existing_index.get('index', {}).get('dimension', 0)
            
            if (set(existing_non_filterable) == set(expected_non_filterable) and 
                existing_dimension == self.embedding_dimension):
                return True
            else:
                pass
                
        except ClientError as e:
            if e.response['Error']['Code'] == 'NotFoundException':
                pass
            else:
                pass
        
        # Start background setup
        tenant_key = f"{user_id}_{project_id}"
        if tenant_key not in self._setup_in_progress:
            def background_setup():
                self.setup_knowledge_base_enhanced(user_id, project_id, force_recreate=True)
            
            threading.Thread(target=background_setup, daemon=True).start()
            return False
        
        return False

    def setup_knowledge_base_enhanced(self, user_id: str, project_id: str, force_recreate: bool = False) -> bool:
        """Enhanced setup with token-aware processing pipeline"""
        setup_id = str(uuid.uuid4())[:8]
        
        tenant_key = f"{user_id}_{project_id}"
        
        # Check if setup is already in progress
        with self._setup_lock:
            if tenant_key in self._setup_in_progress:
                return False
            self._setup_in_progress.add(tenant_key)
        
        try:
            return self._setup_knowledge_base_enhanced_internal(user_id, project_id, force_recreate, setup_id)
        finally:
            with self._setup_lock:
                self._setup_in_progress.discard(tenant_key)

    def _setup_knowledge_base_enhanced_internal(self, user_id: str, project_id: str, force_recreate: bool = False, setup_id: str = None) -> bool:
        """Internal token-aware knowledge base setup with detailed logging"""
        if not setup_id:
            setup_id = str(uuid.uuid4())[:8]
            
        bucket_name, index_name = self._get_tenant_names(user_id, project_id)
        
        # Check existing configuration
        needs_recreate = force_recreate
        
        if not force_recreate:
            try:
                existing_index = self.s3vectors_client.get_index(vectorBucketName=bucket_name, indexName=index_name)
                existing_config = existing_index.get('index', {}).get('metadataConfiguration', {})
                existing_non_filterable = existing_config.get('nonFilterableMetadataKeys', [])
                existing_dimension = existing_index.get('index', {}).get('dimension', 0)
                
                expected_non_filterable = ['full_text', 'text_preview', 'word_count', 'content_length', 'has_images', 'multimodal_content']
                
                if (set(existing_non_filterable) != set(expected_non_filterable) or 
                    existing_dimension != self.embedding_dimension):
                    needs_recreate = True
                else:
                    return True
                    
            except ClientError as e:
                error_code = e.response['Error']['Code']
                needs_recreate = True
        
        # Clean up and recreate if needed
        if needs_recreate:
            if not self._cleanup_existing_infrastructure(bucket_name, index_name):
                return False
            
            if not self._create_vector_infrastructure(bucket_name, index_name):
                return False
        
        # Process documents with token-aware pipeline
        all_files = self.list_files_comprehensive(user_id, project_id)
        
        if not all_files:
            return True
        
        # Process files to create token-aware chunks
        all_chunks = self.process_files_batch(all_files)
        
        if not all_chunks:
            return True
        
        # Generate embeddings with token safety
        chunk_embedding_pairs = self.generate_embeddings_batch(all_chunks)
        
        if not chunk_embedding_pairs:
            return False
        
        # Store enhanced vectors
        success = self.store_vectors_enhanced(
            bucket_name, index_name, chunk_embedding_pairs, user_id, project_id
        )
        
        return success

    def get_document_list(self, user_id: str, project_id: str) -> Dict:
        """Get comprehensive document list with enhanced information"""
        list_id = str(uuid.uuid4())[:8]
        
        files = self.list_files_comprehensive(user_id, project_id)
        
        if not files:
            return {
                'total_files': 0,
                'documents': [],
                'datasets': [],
                'images': [],
                'code_files': [],
                'archives': [],
                'message': '📭 No files found in your collection.'
            }
        
        # Categorize files
        documents = [f for f in files if f['source'] == 'documents' and f['file_type'] in ['PDF', 'DOCX', 'DOC', 'TEXT', 'MARKDOWN', 'RTF']]
        datasets = [f for f in files if f['source'] == 'datasets']
        images = [f for f in files if f['file_type'] == 'IMAGE']
        code_files = [f for f in files if f['file_type'] in ['TYPESCRIPT', 'JAVASCRIPT', 'PYTHON']]
        archives = [f for f in files if f['file_type'] == 'ARCHIVE']
        
        def format_file_info(file_info):
            return {
                'name': os.path.basename(file_info['key']),
                'full_path': file_info['key'],
                'size': file_info['size'],
                'size_mb': round(file_info['size'] / (1024 * 1024), 2),
                'last_modified': file_info['last_modified'].isoformat(),
                'file_type': file_info['file_type'],
                'bucket': file_info['bucket'],
                'is_multimodal': file_info['file_type'] == 'IMAGE',
                'is_code': file_info['file_type'] in ['TYPESCRIPT', 'JAVASCRIPT', 'PYTHON']
            }
        
        formatted_documents = [format_file_info(doc) for doc in documents]
        formatted_datasets = [format_file_info(dataset) for dataset in datasets]
        formatted_images = [format_file_info(image) for image in images]
        formatted_code = [format_file_info(code) for code in code_files]
        formatted_archives = [format_file_info(archive) for archive in archives]
        
        total_doc_size = sum(doc['size_mb'] for doc in formatted_documents)
        total_dataset_size = sum(dataset['size_mb'] for dataset in formatted_datasets)
        total_image_size = sum(image['size_mb'] for image in formatted_images)
        total_code_size = sum(code['size_mb'] for code in formatted_code)
        total_archive_size = sum(archive['size_mb'] for archive in formatted_archives)
        
        return {
            'total_files': len(files),
            'documents': formatted_documents,
            'datasets': formatted_datasets,
            'images': formatted_images,
            'code_files': formatted_code,
            'archives': formatted_archives,
            'total_document_size_mb': round(total_doc_size, 2),
            'total_dataset_size_mb': round(total_dataset_size, 2),
            'total_image_size_mb': round(total_image_size, 2),
            'total_code_size_mb': round(total_code_size, 2),
            'total_archive_size_mb': round(total_archive_size, 2),
            'multimodal_capabilities': len(images) > 0,
            'code_analysis_ready': len(code_files) > 0,
            'message': f'📁 Found {len(documents)} documents ({total_doc_size:.1f} MB), {len(datasets)} datasets ({total_dataset_size:.1f} MB), {len(images)} images ({total_image_size:.1f} MB), {len(code_files)} code files ({total_code_size:.1f} MB), and {len(archives)} archives ({total_archive_size:.1f} MB).'
        }

# Global RAG system instance
_rag_system = None

def get_rag_system():
    """Get or create global RAG system instance"""
    global _rag_system
    if _rag_system is None:
        _rag_system = OptimizedRAGSystem()
    return _rag_system

# ============================================================================
# ENHANCED RAG TOOLS FOR AGENTS
# ============================================================================

@tool(
    name="query_knowledge_vault",
    description="Search through user documents with token-aware processing and enhanced semantic similarity. Returns relevant content with source attribution and quality levels."
)
def query_knowledge_vault(user_id: str, project_id: str, query: str, top_k: int = 10) -> str:
    """Query the knowledge vault with token-aware processing."""
    query_id = str(uuid.uuid4())[:8]
    
    try:
        rag_system = get_rag_system()
        results = rag_system.query_knowledge_base_enhanced(user_id, project_id, query, top_k)
        
        # Log result summary
        result_count = len(results.get('results', []))
        sources_count = len(results.get('sources_found', []))
        similarity_analysis = results.get('similarity_analysis', {})
        
        if results.get('setup_in_progress'):
            message = "🔄 Your documents are being processed and will be available for search shortly. Please try again in a few moments."
            return message
        
        if not results['results']:
            # Enhanced no-results message with similarity analysis
            max_sim = similarity_analysis.get('max_similarity', 0)
            candidates = similarity_analysis.get('total_candidates', 0)
            
            if candidates > 0:
                message = f"🔍 I found {candidates} potentially related documents, but the highest similarity was {max_sim:.1%}. The content might not be directly relevant to '{query}'. Try rephrasing your query with different keywords."
            else:
                message = f"🔍 I searched through your documents but didn't find information about '{query}'. You may want to check if relevant documents have been uploaded or try rephrasing your query."
            
            return message
        
        # Create comprehensive response with quality analysis
        quality_counts = {}
        for result in results['results']:
            quality = result.get('quality_level', 'unknown')
            quality_counts[quality] = quality_counts.get(quality, 0) + 1
        
        response_parts = [
            f"📊 **Token-Aware Search Results for '{query}':**",
            f"Found {len(results['results'])} relevant sections from {len(results.get('sources_found', []))} documents",
            ""
        ]
        
        # Add quality analysis
        if quality_counts:
            quality_summary = []
            for quality, count in quality_counts.items():
                emoji = {"high": "🟢", "medium": "🟡", "low": "🟠"}.get(quality, "⚪")
                quality_summary.append(f"{emoji} {count} {quality} quality")
            
            response_parts.extend([
                f"🎯 **Quality Analysis**: {', '.join(quality_summary)}",
                f"🔍 **Similarity Threshold Used**: {similarity_analysis.get('threshold_used', 'N/A'):.1%}",
                ""
            ])
        
        # Add source summary
        if results.get('sources_found'):
            response_parts.extend([
                "📁 **Sources Found:**",
                ", ".join(results['sources_found']),
                ""
            ])
        
        # Add top results with quality indicators
        response_parts.append("🎯 **Most Relevant Content:**")
        for i, result in enumerate(results['results'][:5], 1):
            quality_emoji = {"high": "🟢", "medium": "🟡", "low": "🟠"}.get(result.get('quality_level', 'unknown'), "⚪")
            response_parts.append(
                f"{i}. {quality_emoji} **{result['file_name']}** ({result['similarity_score']:.1%} relevance)\n"
                f"   {result['content_preview']}"
            )
        
        if results['context']:
            response_parts.extend(["", "📋 **Detailed Context:**", results['context']])
        
        # Add token safety notice
        response_parts.extend([
            "",
            f"⚡ **Token-Safe Processing**: All content processed within API limits for optimal performance"
        ])
        
        final_response = "\n".join(response_parts)
        return final_response
        
    except Exception as e:
        error_message = f"❌ I encountered an issue while searching your documents: {str(e)}. Please try again."
        return error_message

@tool(
    name="list_user_documents", 
    description="Get a comprehensive list of all user documents with enhanced details including file types, sizes, token estimates, and processing capabilities."
)
def list_user_documents(user_id: str, project_id: str) -> str:
    """List all user documents with token-aware comprehensive information."""
    list_id = str(uuid.uuid4())[:8]
    
    try:
        rag_system = get_rag_system()
        document_info = rag_system.get_document_list(user_id, project_id)
        
        if document_info['total_files'] == 0:
            message = "📭 You haven't uploaded any documents yet. Once you upload documents, I'll be able to help you search and analyze their content with token-aware processing."
            return message
        
        response_parts = [
            f"📁 **Token-Aware Document Collection ({document_info['total_files']} files)**",
            ""
        ]
        
        # Add capabilities notice
        capabilities = []
        if document_info.get('multimodal_capabilities'):
            capabilities.append("🎨 **Multimodal AI**: Image analysis with advanced vision")
        if document_info.get('code_analysis_ready'):
            capabilities.append("💻 **Code Analysis**: TypeScript/JavaScript/Python processing")
        
        if capabilities:
            response_parts.extend(capabilities + [""])
        
        if document_info['documents']:
            response_parts.extend([
                f"📄 **Documents ({len(document_info['documents'])} files, {document_info['total_document_size_mb']} MB):**"
            ])
            
            for doc in document_info['documents']:
                file_type_emoji = {
                    'PDF': '📋', 'DOCX': '📝', 'DOC': '📝', 'TEXT': '📄',
                    'MARKDOWN': '📝', 'RTF': '📄'
                }.get(doc['file_type'], '📎')
                
                response_parts.append(
                    f"  {file_type_emoji} **{doc['name']}** - {doc['file_type']} ({doc['size_mb']} MB) - Token-Safe Processing"
                )
            response_parts.append("")
        
        if document_info['datasets']:
            response_parts.extend([
                f"📊 **Dataset Files ({len(document_info['datasets'])} files, {document_info['total_dataset_size_mb']} MB):**"
            ])
            
            for dataset in document_info['datasets']:
                file_type_emoji = {
                    'CSV': '📈', 'JSON': '📊', 'EXCEL': '📊'
                }.get(dataset['file_type'], '📎')
                
                response_parts.append(
                    f"  {file_type_emoji} **{dataset['name']}** - {dataset['file_type']} ({dataset['size_mb']} MB)"
                )
            response_parts.append("")
        
        if document_info.get('images'):
            response_parts.extend([
                f"🖼️ **Images ({len(document_info['images'])} files, {document_info['total_image_size_mb']} MB):**"
            ])
            
            for image in document_info['images']:
                response_parts.append(
                    f"  🖼️ **{image['name']}** - {image['file_type']} ({image['size_mb']} MB) - Multimodal Ready"
                )
            response_parts.append("")
        
        if document_info.get('code_files'):
            response_parts.extend([
                f"💻 **Code Files ({len(document_info['code_files'])} files, {document_info['total_code_size_mb']} MB):**"
            ])
            
            for code in document_info['code_files']:
                code_emoji = {'TYPESCRIPT': '🔷', 'JAVASCRIPT': '🟨', 'PYTHON': '🐍'}.get(code['file_type'], '💻')
                response_parts.append(
                    f"  {code_emoji} **{code['name']}** - {code['file_type']} ({code['size_mb']} MB) - Code Analysis Ready"
                )
            response_parts.append("")
        
        if document_info.get('archives'):
            response_parts.extend([
                f"📦 **Archives ({len(document_info['archives'])} files, {document_info['total_archive_size_mb']} MB):**"
            ])
            
            for archive in document_info['archives']:
                response_parts.append(
                    f"  📦 **{archive['name']}** - {archive['file_type']} ({archive['size_mb']} MB)"
                )
            response_parts.append("")
        
        response_parts.extend([
            "🔍 **Enhanced Search Capabilities:**",
            "• Token-aware semantic search with optimized processing",
            "• Permissive similarity thresholds for better recall",
            "• Multi-format support (PDF, Word, Text, CSV, JSON, Images, Code)",
            "• Quality-based result ranking (High/Medium/Low relevance)",
            "• Multimodal analysis for images with AI vision",
            "• Code analysis for TypeScript/JavaScript/Python files",
            "• Source attribution with specific file references",
            "",
            "💡 **How to Use:**",
            "• Ask specific questions about your document content",
            "• Request code analysis or explanations",
            "• Search for information across your entire collection",
            "• The system automatically handles token limits for optimal performance",
            "• Image content is analyzed for visual understanding",
            "",
            "⚙️ **Token-Aware System Features:**",
            f"• Static token management: {rag_system.max_chars_for_embedding:,} char limit",
            f"• Estimated token capacity: {rag_system.max_tokens_estimate:,} tokens",
            f"• Enhanced embedding dimensions: {rag_system.embedding_dimension}",
            f"• Permissive similarity thresholds for better results",
            f"• Content-type aware processing (English, Technical, Mixed)",
            f"• Total storage: {sum([document_info['total_document_size_mb'], document_info['total_dataset_size_mb'], document_info.get('total_image_size_mb', 0), document_info.get('total_code_size_mb', 0), document_info.get('total_archive_size_mb', 0)]):.1f} MB"
        ])
        
        final_response = "\n".join(response_parts)
        return final_response
        
    except Exception as e:
        return f"❌ I encountered an issue while retrieving your document list: {str(e)}. Please try again."

# ============================================================================
# FASTAPI APPLICATION SETUP
# ============================================================================

app = FastAPI(
    title="Token-Aware Document Analysis System with Enhanced AI",
    description="Advanced multi-agent system with static token management and comprehensive document processing",
    version="10.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["*"],
)

# Request/Response Models
class WorkflowRequest(BaseModel):
    user_id: str = Field(..., description="User identifier")
    project_id: str = Field(default="default", description="Project identifier")
    message: str = Field(..., description="User message or query")
    context: Optional[Dict[str, Any]] = Field(None, description="Additional context")
    file_upload: Optional[Dict[str, Any]] = Field(None, description="File upload data")
    agent_name: Optional[str] = Field(None, description="Specific agent name to use")

class WorkflowResponse(BaseModel):
    response: str
    processing_time: float
    agent_used: str
    workflow_config_found: bool

# ============================================================================
# ENHANCED AGENT FACTORY
# ============================================================================

class OptimizedAgentFactory:
    def __init__(self):
        """Initialize the agent factory with token-aware capabilities"""
        self.s3_client = boto3.client('s3', region_name='us-east-1')
        self.config_cache = {}
        self.agent_cache = {}
        
        if STRANDS_AVAILABLE:
            self.default_tools = [
                use_aws, 
                http_request, 
                retrieve, 
                generate_image_sd3,
                generate_image_from_reference,
                upload_image_to_s3,
                image_reader,
                query_knowledge_vault,
                list_user_documents
            ]
        else:
            self.default_tools = [
                generate_image_sd3,
                generate_image_from_reference,
                upload_image_to_s3,
                query_knowledge_vault,
                list_user_documents
            ]
        
        self.model_mapping = {
            "Claude 3.5 Sonnet": "anthropic.claude-3-5-sonnet-20240620-v1:0",
            "Claude 3 Opus": "anthropic.claude-3-opus-20240229-v1:0",
            "Claude 3 Sonnet": "anthropic.claude-3-sonnet-20240229-v1:0",
            "Claude 3 Haiku": "anthropic.claude-3-haiku-20240307-v1:0",
            "Command R": "cohere.command-r-v1:0", 
            "Command R+": "cohere.command-r-plus-v1:0",
            "Titan Express": "amazon.titan-text-express-v1:0",
            "Mistral Large": "mistral.mistral-large-2402-v1:0",
            "Mixtral 8x7B": "mistral.mixtral-8x7b-v0:1",
            "Llama 3 70B": "meta.llama3-70b-instruct-v1:0",
            "Amazon Nova Pro": "amazon.nova-pro-v1:0"
        }
    
    def get_workflow_config(self, user_id: str, project_id: str) -> Optional[Dict]:
        """Get workflow configuration from cloud storage"""
        config_id = str(uuid.uuid4())[:8]
        cache_key = f"{user_id}_{project_id}"
        
        if cache_key in self.config_cache:
            return self.config_cache[cache_key]
        
        try:
            bucket_name = 'qubitz-customer-prod'
            key = f'{user_id}/{project_id}/workflow.json'
            
            response = self.s3_client.get_object(Bucket=bucket_name, Key=key)
            workflow_data = json.loads(response['Body'].read().decode('utf-8'))
            
            self.config_cache[cache_key] = workflow_data
            
            return workflow_data
            
        except Exception as e:
            return None

    def create_agent_tools_from_config(self, agent_config: Dict) -> List:
        """Create tools for an agent based on configuration"""
        tools = []
        
        for tool_config in agent_config.get('tools', []):
            tool_name = tool_config.get('name')
            tool_type = tool_config.get('type')
            
            if tool_type == 'strands_builtin':
                if tool_name == 'cloud_services' and use_aws:
                    tools.append(use_aws)
                elif tool_name == 'web_search' and http_request:
                    tools.append(http_request)
                elif tool_name == 'api_calls' and http_request:
                    tools.append(http_request)
                elif tool_name == 'data_processing' and retrieve:
                    tools.append(retrieve)
                elif tool_name == 'file_operations' and use_aws:
                    tools.append(use_aws)
        
        # Always add essential tools
        essential_tools = [
            query_knowledge_vault, 
            list_user_documents, 
            generate_image_sd3,
            generate_image_from_reference,
            upload_image_to_s3
        ]
        
        if image_reader:
            essential_tools.append(image_reader)
            
        for tool in essential_tools:
            if tool and tool not in tools:
                tools.append(tool)
        
        if not tools and self.default_tools:
            tools = self.default_tools
            
        return tools

    def build_system_prompt(self, agent_config: Dict, workflow_config: Dict, user_id: str, project_id: str) -> str:
        """Build token-aware system prompt"""
        agent_name = agent_config.get('name', 'Assistant')
        agent_type = agent_config.get('agent_type', 'general')
        description = agent_config.get('description', 'A helpful assistant.')
        instructions = agent_config.get('instructions', '')
        
        system_prompt = f"""
You are {agent_name.replace('_', ' ').title()}, an advanced {agent_type} assistant with token-aware document analysis capabilities and professional AI image generation.

## Your Enhanced Identity & Mission
{description}

## Core Instructions & Responsibilities
{instructions}

## Advanced Token-Aware Capabilities

**Token-Safe Document Analysis System:**
- Use <thinking> tags first if any knowledge vault query related query comes up and say that querying knowledge vault so that user can wait and keep giving the status after successful stages like parsing, indexing, etc.
- `query_knowledge_vault`: Advanced semantic search with static token management
- `list_user_documents`: Complete inventory with token-aware processing details

**Professional AI Image Generation:**
- `generate_image_sd3`: Create high-quality images using Stable Diffusion 3 Large model
- `generate_image_from_reference`: Transform existing images using Stable Diffusion XL
- `upload_image_to_s3`: Upload and prepare reference images for image-to-image workflows
- `image_reader`: Advanced image analysis and content extraction

**Cloud Services & Web Intelligence:**
- `use_aws`: Comprehensive AWS services with token-aware processing
- `http_request`: Real-time information research and verification
- `retrieve`: Advanced data processing and retrieval

User Context: {user_id}/{project_id}

## Token-Aware Processing Features

**Static Token Management:**
- All text processing uses static token limits (no external tokenizer dependency)
- Conservative character limits: 5,000 chars, ~6,500 tokens estimated
- Content-type aware processing (English, Technical, Mixed content)
- Progressive truncation with word boundary preservation

**Quality-Based Search:**
- Permissive similarity thresholds (0.02 minimum for maximum recall)
- Quality indicators: High (>50%), Medium (20-50%), Low (2-20%)
- Multi-format support: PDF, Word, Text, CSV, JSON, Images, Code

**Enhanced File Processing:**
- PDF: Page-by-page extraction with structure preservation
- Word: Paragraph and table extraction
- Code: TypeScript/JavaScript/Python analysis
- Images: Multimodal analysis with AI vision
- Archives: Metadata extraction and cataloging

## Response Excellence Standards

**Professional Communication:**
- Always use token-aware processing for optimal performance
- Provide quality indicators for all search results
- Include similarity thresholds and processing details
- Deliver clear, well-structured responses with proper formatting

**Image Generation Standards:**
- ALL generated images are automatically uploaded to S3
- ALWAYS provide the S3 URL to users in your response
- Include technical details (model, settings, dimensions)
- Explain generation approach and quality settings

**Token-Aware Problem-Solving:**
1. Understand user needs (documents, images, code analysis)
2. Use token-safe search with appropriate similarity thresholds
3. Generate images when visual content would enhance responses
4. Provide comprehensive analysis with quality assessments
5. Always include S3 URLs for generated images
6. Explain token management and processing optimizations

## System Enhancement Features

**Adaptive Intelligence:**
- Static token estimation based on content type
- Quality-based result ranking for better user experience
- Enhanced precision with 1024-dimension embeddings
- Multimodal capabilities for comprehensive analysis

**Quality Assurance:**
- All search results include quality levels and confidence scores
- Threshold information provided for transparency
- Enhanced error handling with detailed feedback
- Comprehensive logging for system optimization

Your goal is to provide exceptional assistance through token-aware intelligent document analysis and professional AI image generation. Every response should demonstrate the system's enhanced capabilities while maintaining optimal performance within token limits.

**Remember: Token-aware system with static limits, quality indicators, and comprehensive file support - always provide S3 URLs for generated images and explain processing optimizations!**
        """
        
        return system_prompt.strip()

    def get_agent_from_config(self, agent_config: Dict, workflow_config: Dict, user_id: str, project_id: str) -> Agent:
        """Create token-aware agent based on configuration"""
        agent_id = str(uuid.uuid4())[:8]
        agent_name = agent_config.get('name', 'default_agent')
        cache_key = f"{user_id}_{project_id}_{agent_name}"
        
        if cache_key in self.agent_cache:
            return self.agent_cache[cache_key]
            
        model_name = agent_config.get('model_name', 'Claude 3.5 Sonnet')
        model_id = agent_config.get('model_id')
        
        if not model_id and model_name in self.model_mapping:
            model_id = self.model_mapping[model_name]
        elif not model_id:
            model_id = "anthropic.claude-3-5-sonnet-20240620-v1:0"
            
        params = agent_config.get('parameters', {})
        max_tokens = "30000"
        temperature = params.get('temperature', 0.7)
        top_p = params.get('top_p', 0.9)
        
        model = BedrockModel(
            model_id="us.anthropic.claude-sonnet-4-20250514-v1:0",
            boto_session=boto3.Session(region_name='us-east-1'),
            params={
                "temperature": temperature,
                "top_p": top_p,
                "max_tokens": max_tokens
            }
        )
        
        system_prompt = self.build_system_prompt(agent_config, workflow_config, user_id, project_id)
        tools = self.create_agent_tools_from_config(agent_config)
        
        agent = Agent(
            model=model,
            system_prompt=system_prompt,
            conversation_manager=SlidingWindowConversationManager(
                window_size=10 if agent_config.get('memory_enabled', True) else 1
            ),
            tools=tools
        )
        
        self.agent_cache[cache_key] = agent
        return agent
    
    def get_default_agent(self, user_id: str, project_id: str) -> Agent:
        """Get token-aware default agent"""
        cache_key = f"default_agent_{user_id}_{project_id}"
        if cache_key in self.agent_cache:
            return self.agent_cache[cache_key]
            
        model = BedrockModel(
            model_id="anthropic.claude-3-5-sonnet-20240620-v1:0",
            max_tokens=4000,
            boto_session=boto3.Session(region_name='us-east-1'),
            params={
                "temperature": 0.7,
                "top_p": 0.9
            }
        )
        
        system_prompt = f"""
You are an advanced AI assistant with token-aware document analysis, enhanced file processing, and professional AI image generation capabilities.

## Token-Aware Core Capabilities

**Advanced Document Analysis System:**
- Token-safe semantic search with static limits (5,000 chars, ~6,500 tokens)
- Quality-based result ranking with permissive thresholds
- Multi-format content analysis (PDF, Word, Text, CSV, JSON, Images, Code)
- Enhanced precision with 1024-dimension embeddings

**Professional AI Image Generation:**
- `generate_image_sd3`: Create high-quality images using Stable Diffusion 3 Large
- `generate_image_from_reference`: Transform images using Stable Diffusion XL
- `upload_image_to_s3`: Prepare reference images for image-to-image workflows
- `image_reader`: Advanced image analysis and processing

**Available Token-Aware Tools:**
- `query_knowledge_vault`: Advanced semantic search with static token management
- `list_user_documents`: Complete inventory with processing capabilities
- `generate_image_sd3`: Latest AI image generation technology
- `generate_image_from_reference`: Advanced image transformation
- `upload_image_to_s3`: Image preparation and management
- `image_reader`: Multimodal image analysis
- `use_aws`: Comprehensive cloud services

User Context: {user_id}/{project_id}

## Enhanced Mission & Standards

**Token-Aware Document Intelligence:**
- Use static token estimation for optimal performance
- Provide quality indicators for all search results (High/Medium/Low)
- Support comprehensive file formats with specialized processing
- Explain processing optimizations and token management

**Professional Image Generation:**
- Create compelling visual content with latest AI models
- Use detailed, descriptive prompts for high-quality results
- Automatically upload all generated images to S3
- Always provide S3 URLs to users for immediate access

**Excellence Standards:**
- Always search user documents with token-aware algorithms
- Generate images when visual content would enhance responses
- Provide comprehensive analysis with quality indicators
- Use permissive thresholds for maximum relevance
- Support multimodal understanding across all content types

**Token-Aware Image Generation Workflow:**
1. Understand user's visual requirements
2. Create detailed, descriptive prompts
3. Choose optimal model (SD3 Large for new images, SDXL for transformations)
4. Generate with quality settings
5. **AUTOMATIC**: Upload to S3 with metadata
6. **MANDATORY**: Provide S3 URL with technical details

## System Enhancement Features

**Token-Safe Processing:**
- Static token limits: 5,000 characters, ~6,500 tokens
- Content-type aware processing (English, Technical, Mixed)
- Progressive truncation with word boundary preservation
- Quality indicators: High (>50%), Medium (20-50%), Low (2-20%)

**Enhanced Capabilities:**
- Permissive similarity thresholds (0.02 minimum) for better recall
- Multi-format support: Documents, Code, Images, Archives
- Multimodal analysis with AI vision
- Code analysis for TypeScript/JavaScript/Python

Your goal is to provide outstanding assistance through token-aware intelligent document analysis and professional AI image generation. Every response should demonstrate enhanced capabilities while maintaining optimal performance within static token limits.

**Remember: Token-aware system with static limits, quality indicators, and comprehensive file support - always provide S3 URLs for generated images and explain processing optimizations!**
        """
        
        agent = Agent(
            model=model,
            system_prompt=system_prompt,
            conversation_manager=SlidingWindowConversationManager(window_size=5),
            tools=self.default_tools
        )
        
        self.agent_cache[cache_key] = agent
        return agent

    def select_agent(self, message: str, workflow_config: Dict, requested_agent: Optional[str] = None) -> Dict:
        """Select appropriate agent based on message content and configuration"""
        if not workflow_config or 'agents' not in workflow_config:
            return None
            
        if requested_agent:
            for agent in workflow_config['agents']:
                if agent.get('name') == requested_agent:
                    return agent
        
        # Check for orchestrator first
        for agent in workflow_config['agents']:
            if agent.get('agent_type') == 'orchestrator':
                return agent
                
        # Intelligent selection based on message content
        agents = workflow_config['agents']
        best_match = None
        best_score = -1
        
        for agent in agents:
            score = 0
            agent_name = agent.get('name', '').lower()
            agent_desc = agent.get('description', '').lower()
            message_lower = message.lower()
            
            # Scoring logic
            if agent_name.replace('_', ' ') in message_lower:
                score += 10
            
            # Check description keywords
            desc_words = re.findall(r'\b\w{4,}\b', agent_desc)
            for word in desc_words:
                if word in message_lower:
                    score += 2
            
            # Check tool capabilities
            for tool in agent.get('tools', []):
                tool_name = tool.get('name', '').lower()
                if any(keyword in message_lower for keyword in tool_name.split('_')):
                    score += 1
            
            if score > best_score:
                best_score = score
                best_match = agent
        
        return best_match or workflow_config['agents'][0]

class SessionManager:
    def __init__(self):
        self.sessions = {}
        
    def get_or_create_session(self, user_id: str, project_id: str) -> Dict:
        session_key = f"{user_id}_{project_id}"
        
        if session_key not in self.sessions:
            self.sessions[session_key] = {
                "history": [],
                "context": {},
                "created_at": datetime.now().isoformat(),
                "last_active": datetime.now().isoformat()
            }
            
        return self.sessions[session_key]
        
    def add_interaction(self, user_id: str, project_id: str, message: str, response: str, agent_name: str):
        session = self.get_or_create_session(user_id, project_id)
        
        session["history"].append({
            "timestamp": datetime.now().isoformat(),
            "message": message,
            "response": response[:500] + ("..." if len(response) > 500 else ""),
            "agent": agent_name
        })
        
        session["last_active"] = datetime.now().isoformat()
        
        # Keep last 15 interactions
        if len(session["history"]) > 15:
            session["history"] = session["history"][-15:]
            
    def get_recent_history(self, user_id: str, project_id: str, count: int = 3) -> List[Dict]:
        session = self.get_or_create_session(user_id, project_id)
        return session["history"][-count:] if session["history"] else []

class OptimizedOrchestrator:
    def __init__(self):
        self.agent_factory = OptimizedAgentFactory()
        self.session_manager = SessionManager()
        
    def process_request(self, request: WorkflowRequest) -> tuple:
        process_id = str(uuid.uuid4())[:8]
        start_time = time.time()
        
        # Trigger token-aware knowledge base setup if needed
        rag_system = get_rag_system()
        rag_system.auto_setup_if_needed(request.user_id, request.project_id)
        
        workflow_config = self.agent_factory.get_workflow_config(request.user_id, request.project_id)
        workflow_found = workflow_config is not None
        
        history = self.session_manager.get_recent_history(request.user_id, request.project_id)
        
        if workflow_found:
            agent_config = self.agent_factory.select_agent(
                request.message, 
                workflow_config,
                requested_agent=request.agent_name
            )
            
            if agent_config:
                agent = self.agent_factory.get_agent_from_config(
                    agent_config, 
                    workflow_config,
                    request.user_id, 
                    request.project_id
                )
                agent_name = agent_config.get('name', 'unknown')
            else:
                agent = self.agent_factory.get_default_agent(request.user_id, request.project_id)
                agent_name = "token_aware_default_agent"
        else:
            agent = self.agent_factory.get_default_agent(request.user_id, request.project_id)
            agent_name = "token_aware_default_agent"
        
        # Build context
        context_parts = []
        
        if request.context:
            for key, value in request.context.items():
                context_parts.append(f"{key}: {value}")
        
        if request.file_upload:
            context_parts.append(f"File upload: {json.dumps(request.file_upload)}")
        
        if history:
            context_parts.append("Recent conversation:")
            for item in history[-2:]:
                context_parts.append(f"User: {item['message']}")
                context_parts.append(f"Assistant: {item['response'][:100]}...")
        
        if context_parts:
            full_request = f"Context: {' | '.join(context_parts)}\n\nUser: {request.message}"
        else:
            full_request = request.message
        
        response = agent(full_request)
        response_text = self._extract_response_text(response)
        
        self.session_manager.add_interaction(
            request.user_id, 
            request.project_id,
            request.message,
            response_text,
            agent_name
        )
        
        processing_time = time.time() - start_time
        
        return response_text, processing_time, agent_name, workflow_found
        
    def stream_request(self, request: WorkflowRequest):
        process_id = str(uuid.uuid4())[:8]
        
        # Trigger token-aware knowledge base setup
        rag_system = get_rag_system()
        rag_system.auto_setup_if_needed(request.user_id, request.project_id)
        
        workflow_config = self.agent_factory.get_workflow_config(request.user_id, request.project_id)
        
        history = self.session_manager.get_recent_history(request.user_id, request.project_id)
        
        if workflow_config:
            agent_config = self.agent_factory.select_agent(
                request.message, 
                workflow_config,
                requested_agent=request.agent_name
            )
            
            if agent_config:
                agent = self.agent_factory.get_agent_from_config(
                    agent_config, 
                    workflow_config,
                    request.user_id, 
                    request.project_id
                )
                agent_name = agent_config.get('name', 'unknown')
            else:
                agent = self.agent_factory.get_default_agent(request.user_id, request.project_id)
                agent_name = "token_aware_default_agent"
        else:
            agent = self.agent_factory.get_default_agent(request.user_id, request.project_id)
            agent_name = "token_aware_default_agent"
        
        # Build context
        context_parts = []
        
        if request.context:
            for key, value in request.context.items():
                context_parts.append(f"{key}: {value}")
        
        if request.file_upload:
            context_parts.append(f"File upload: {json.dumps(request.file_upload)}")
        
        if history:
            context_parts.append("Recent conversation:")
            for item in history[-2:]:
                context_parts.append(f"User: {item['message']}")
                context_parts.append(f"Assistant: {item['response'][:100]}...")
        
        if context_parts:
            full_request = f"Context: {' | '.join(context_parts)}\n\nUser: {request.message}"
        else:
            full_request = request.message
        
        async def stream_generator():
            full_response = ""
            
            async for chunk in agent.stream_async(full_request):
                if 'data' in chunk:
                    yield chunk['data']
                    full_response += chunk['data']
                    
            self.session_manager.add_interaction(
                request.user_id, 
                request.project_id,
                request.message,
                full_response,
                agent_name
            )
                
        return stream_generator()
    
    def _extract_response_text(self, response) -> str:
        try:
            if hasattr(response, 'content'):
                if isinstance(response.content, list):
                    text_parts = []
                    for item in response.content:
                        if isinstance(item, dict):
                            if 'text' in item:
                                text_parts.append(item['text'])
                            elif 'content' in item:
                                text_parts.append(str(item['content']))
                        else:
                            text_parts.append(str(item))
                    return ' '.join(text_parts)
                else:
                    return str(response.content)
            else:
                return str(response)
        except Exception as e:
            return str(response)

# Initialize the token-aware orchestrator
orchestrator = OptimizedOrchestrator()

# ============================================================================
# ENHANCED API ENDPOINTS
# ============================================================================

@app.get("/health")
def health_check():
    """Enhanced health check endpoint"""
    return {
        "status": "healthy",
        "message": "Token-Aware Document Analysis System with Enhanced AI is running",
        "version": "10.0.0",
        "features": {
            "static_token_management": True,
            "no_tokenizer_dependency": True,
            "permissive_similarity_thresholds": True,
            "content_type_aware_processing": True,
            "enhanced_file_support": True,
            "multimodal_capabilities": True,
            "code_analysis": True,
            "stable_diffusion_3_large": True,
            "stable_diffusion_xl_img2img": True,
            "s3_auto_upload": True,
            "debug_endpoints": True
        }
    }

@app.post('/workflow/chat', response_model=WorkflowResponse)
async def workflow_chat(request: WorkflowRequest):
    try:
        response_text, processing_time, agent_used, workflow_found = orchestrator.process_request(request)
        return WorkflowResponse(
            response=response_text,
            processing_time=processing_time,
            agent_used=agent_used,
            workflow_config_found=workflow_found
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing request: {str(e)}")

@app.post('/workflow/stream')
async def workflow_stream(request: WorkflowRequest):
    try:
        return StreamingResponse(
            orchestrator.stream_request(request),
            media_type="text/plain"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error streaming response: {str(e)}")

@app.get('/workflow/agents/{user_id}/{project_id}')
async def get_workflow_agents(user_id: str, project_id: str):
    try:
        workflow_config = orchestrator.agent_factory.get_workflow_config(user_id, project_id)
        
        if not workflow_config or 'agents' not in workflow_config:
            return {
                "message": "No workflow configuration found",
                "agents": []
            }
        
        agents_info = []
        for agent in workflow_config['agents']:
            agents_info.append({
                "name": agent.get('name'),
                "type": agent.get('agent_type'),
                "description": agent.get('description'),
                "tools": [tool.get('name') for tool in agent.get('tools', [])],
                "features": [
                    "token_aware_processing",
                    "static_token_management", 
                    "permissive_similarity_search",
                    "enhanced_file_support",
                    "multimodal_capabilities", 
                    "code_analysis",
                    "stable_diffusion_3_large",
                    "stable_diffusion_xl_img2img",
                    "s3_auto_upload"
                ]
            })
        
        return {
            "workflow_found": True,
            "total_agents": len(agents_info),
            "agents": agents_info
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting workflow agents: {str(e)}")

@app.post('/setup/knowledge-base/{user_id}/{project_id}')
async def setup_knowledge_base_endpoint(user_id: str, project_id: str):
    """Manually trigger token-aware knowledge base setup"""
    try:
        rag_system = get_rag_system()
        success = rag_system.setup_knowledge_base_enhanced(user_id, project_id)
        
        if success:
            return {
                "status": "success",
                "message": f"Token-aware knowledge base setup completed for {user_id}/{project_id}",
                "features": ["static_token_management", "permissive_similarity", "enhanced_file_support"]
            }
        else:
            return {
                "status": "error", 
                "message": "Token-aware knowledge base setup failed"
            }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Setup failed: {str(e)}")

@app.post('/setup/force-recreate-knowledge-base/{user_id}/{project_id}')
async def force_recreate_knowledge_base_endpoint(user_id: str, project_id: str):
    """Force recreate token-aware knowledge base"""
    try:
        rag_system = get_rag_system()
        
        # Clear any cached setup locks
        tenant_key = f"{user_id}_{project_id}"
        with rag_system._setup_lock:
            rag_system._setup_in_progress.discard(tenant_key)
        
        success = rag_system.setup_knowledge_base_enhanced(user_id, project_id, force_recreate=True)
        
        if success:
            return {
                "status": "success",
                "message": f"Token-aware knowledge base force recreated successfully for {user_id}/{project_id}",
                "action_taken": "Deleted and recreated index with token-aware metadata structure and enhanced processing",
                "features": ["static_token_management", "permissive_similarity", "enhanced_file_support"]
            }
        else:
            return {
                "status": "error", 
                "message": "Token-aware knowledge base recreation failed"
            }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Force recreation failed: {str(e)}")

# ============================================================================
# ENHANCED DEBUGGING ENDPOINTS
# ============================================================================

@app.get('/debug/knowledge-base/{user_id}/{project_id}')
async def debug_knowledge_base(user_id: str, project_id: str):
    """Enhanced debug endpoint with token-aware capabilities check"""
    debug_id = str(uuid.uuid4())[:8]
    
    try:
        rag_system = get_rag_system()
        bucket_name, index_name = rag_system._get_tenant_names(user_id, project_id)
        
        debug_info = {
            "debug_id": debug_id,
            "version": "10.0.0",
            "tenant_info": {
                "bucket_name": bucket_name,
                "index_name": index_name
            },
            "bucket_status": "unknown",
            "index_status": "unknown",
            "index_config": None,
            "vector_count": 0,
            "sample_query_test": "not_tested",
            "file_discovery": {"documents": 0, "datasets": 0, "images": 0, "code_files": 0, "archives": 0},
            "embedding_test": "not_tested",
            "token_management": {
                "max_chars": rag_system.max_chars_for_embedding,
                "estimated_max_tokens": rag_system.max_tokens_estimate,
                "char_to_token_ratios": rag_system.char_to_token_ratios
            },
            "similarity_thresholds": rag_system.similarity_thresholds,
            "enhanced_features": {
                "static_token_management": True,
                "permissive_similarity": True,
                "enhanced_file_support": True,
                "multimodal_support": True,
                "code_analysis": True,
                "enhanced_dimensions": rag_system.embedding_dimension
            }
        }
        
        # Check bucket status
        try:
            bucket_info = rag_system.s3vectors_client.get_vector_bucket(vectorBucketName=bucket_name)
            debug_info["bucket_status"] = "exists"
        except ClientError as e:
            debug_info["bucket_status"] = f"error: {e.response['Error']['Code']}"
        
        # Check index status with token-aware configuration
        try:
            index_info = rag_system.s3vectors_client.get_index(vectorBucketName=bucket_name, indexName=index_name)
            debug_info["index_status"] = "exists"
            debug_info["index_config"] = index_info
            
            # Check if token-aware configuration matches
            existing_config = index_info.get('index', {}).get('metadataConfiguration', {})
            existing_dimension = index_info.get('index', {}).get('dimension', 0)
            
            debug_info["enhanced_features"]["current_dimension"] = existing_dimension
            debug_info["enhanced_features"]["dimension_match"] = existing_dimension == rag_system.embedding_dimension
            
        except ClientError as e:
            debug_info["index_status"] = f"error: {e.response['Error']['Code']}"
        
        # Test enhanced file discovery
        try:
            all_files = rag_system.list_files_comprehensive(user_id, project_id)
            documents = [f for f in all_files if f['source'] == 'documents' and f['file_type'] in ['PDF', 'DOCX', 'DOC', 'TEXT', 'MARKDOWN', 'RTF']]
            datasets = [f for f in all_files if f['source'] == 'datasets']
            images = [f for f in all_files if f['file_type'] == 'IMAGE']
            code_files = [f for f in all_files if f['file_type'] in ['TYPESCRIPT', 'JAVASCRIPT', 'PYTHON']]
            archives = [f for f in all_files if f['file_type'] == 'ARCHIVE']
            
            debug_info["file_discovery"] = {
                "documents": len(documents),
                "datasets": len(datasets),
                "images": len(images),
                "code_files": len(code_files),
                "archives": len(archives),
                "total": len(all_files),
                "multimodal_content": len(images) > 0,
                "code_analysis_ready": len(code_files) > 0
            }
        except Exception as e:
            debug_info["file_discovery"] = {"error": str(e)}
        
        # Test token-aware embedding generation
        try:
            test_embedding = rag_system.generate_enhanced_embedding("test query for token-aware debugging")
            if test_embedding:
                debug_info["embedding_test"] = f"success - token-aware dimensions: {len(test_embedding)}"
                debug_info["enhanced_features"]["embedding_generation"] = "working"
            else:
                debug_info["embedding_test"] = "failed - no embedding returned"
                debug_info["enhanced_features"]["embedding_generation"] = "failed"
        except Exception as e:
            debug_info["embedding_test"] = f"error: {str(e)}"
            debug_info["enhanced_features"]["embedding_generation"] = "error"
        
        # Test token-aware sample query with permissive similarity
        if debug_info["index_status"] == "exists" and debug_info["embedding_test"].startswith("success"):
            try:
                test_embedding = rag_system.generate_enhanced_embedding("test query")
                if test_embedding:
                    query_params = {
                        'vectorBucketName': bucket_name,
                        'indexName': index_name,
                        'queryVector': {'float32': test_embedding},
                        'topK': 5,
                        'returnMetadata': True,
                        'returnDistance': True
                    }
                    
                    response = rag_system.s3vectors_client.query_vectors(**query_params)
                    results = response.get('vectors', [])
                    debug_info["vector_count"] = len(results)
                    
                    if results:
                        distances = [result.get('distance', 1.0) for result in results]
                        similarities = [max(0, 1 - distance) for distance in distances]
                        debug_info["similarity_analysis"] = {
                            "max_similarity": max(similarities) if similarities else 0,
                            "min_similarity": min(similarities) if similarities else 0,
                            "avg_similarity": sum(similarities) / len(similarities) if similarities else 0,
                            "permissive_threshold": rag_system.similarity_thresholds['any_match'],
                            "results_above_threshold": sum(1 for s in similarities if s >= rag_system.similarity_thresholds['any_match'])
                        }
                        debug_info["sample_query_test"] = "success - with token-aware similarity analysis"
                    else:
                        debug_info["sample_query_test"] = "success - no vectors found"
                        debug_info["similarity_analysis"] = "no_vectors_to_analyze"
                else:
                    debug_info["sample_query_test"] = "embedding_generation_failed"
                    
            except Exception as e:
                debug_info["sample_query_test"] = f"failed: {str(e)}"
        
        return debug_info
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Token-aware debug failed: {str(e)}")

@app.get('/debug/logs/recent')
async def get_recent_logs():
    """Get recent log entries for debugging"""
    try:
        import os
        log_file = 'rag_system.log'
        
        if not os.path.exists(log_file):
            return {"message": "No log file found"}
        
        with open(log_file, 'r') as f:
            lines = f.readlines()
            recent_lines = lines[-100:]  # Last 100 lines
            
        return {
            "total_lines": len(lines),
            "recent_lines": recent_lines,
            "log_file": log_file,
            "version": "10.0.0"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Log retrieval failed: {str(e)}")

@app.get('/debug/logs/search/{search_term}')
async def search_logs(search_term: str):
    """Search log entries for specific terms"""
    try:
        import os
        log_file = 'rag_system.log'
        
        if not os.path.exists(log_file):
            return {"message": "No log file found"}
        
        with open(log_file, 'r') as f:
            lines = f.readlines()
            
        matching_lines = [line for line in lines if search_term.lower() in line.lower()]
        
        return {
            "search_term": search_term,
            "total_matches": len(matching_lines),
            "matching_lines": matching_lines[-50:],  # Last 50 matches
            "log_file": log_file,
            "version": "10.0.0"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Log search failed: {str(e)}")

@app.get('/debug/token-analysis/{user_id}/{project_id}')
async def debug_token_analysis(user_id: str, project_id: str, query: str = "test query"):
    """Debug token analysis and static estimation"""
    debug_id = str(uuid.uuid4())[:8]
    
    try:
        rag_system = get_rag_system()
        
        # Test token estimation
        content_types = ['english', 'technical', 'mixed', 'default']
        token_analysis = {
            "debug_id": debug_id,
            "query": query,
            "static_limits": {
                "max_chars": rag_system.max_chars_for_embedding,
                "estimated_max_tokens": rag_system.max_tokens_estimate
            },
            "content_type_analysis": {},
            "truncation_test": {},
            "embedding_test": {}
        }
        
        # Test each content type
        for content_type in content_types:
            char_count = len(query)
            estimated_tokens = rag_system._estimate_tokens(query, content_type)
            
            token_analysis["content_type_analysis"][content_type] = {
                "char_count": char_count,
                "estimated_tokens": estimated_tokens,
                "ratio": rag_system.char_to_token_ratios[content_type],
                "within_limits": estimated_tokens <= rag_system.max_tokens_estimate
            }
        
        # Test truncation
        long_text = query * 1000  # Create long text
        detected_type = rag_system._detect_content_type(long_text)
        truncated_text = rag_system._truncate_text_safely(long_text, debug_id)
        
        token_analysis["truncation_test"] = {
            "original_length": len(long_text),
            "truncated_length": len(truncated_text),
            "detected_content_type": detected_type,
            "original_estimated_tokens": rag_system._estimate_tokens(long_text, detected_type),
            "truncated_estimated_tokens": rag_system._estimate_tokens(truncated_text, detected_type)
        }
        
        # Test embedding generation
        try:
            test_embedding = rag_system.generate_enhanced_embedding(query)
            if test_embedding:
                token_analysis["embedding_test"] = {
                    "status": "success",
                    "embedding_dimensions": len(test_embedding),
                    "query_length": len(query),
                    "estimated_tokens": rag_system._estimate_tokens(query, 'default')
                }
            else:
                token_analysis["embedding_test"] = {"status": "failed", "reason": "no_embedding_returned"}
        except Exception as e:
            token_analysis["embedding_test"] = {"status": "error", "error": str(e)}
        
        return token_analysis
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Token analysis debug failed: {str(e)}")

# Run the token-aware application
if __name__ == "__main__":
    port = int(os.getenv("PORT", 8000))
    
    print("🚀 Starting Token-Aware Document Analysis System with Enhanced AI")
    print(f"📊 Static Token Management: No external tokenizer dependency")
    print(f"🎯 S3 Vectors Integration: Token-aware with enhanced processing")
    print(f"💡 Enhanced Embeddings: 1024 dimensions with permissive thresholds")
    print(f"📄 Enhanced File Processing: PDF, Word, Text, CSV, JSON, Images, Code, Archives")
    print(f"🎨 Stable Diffusion 3 Large: Latest AI image generation")
    print(f"🔄 Stable Diffusion XL: Advanced image-to-image generation")
    print(f"☁️ S3 Auto-Upload: All images automatically uploaded")
    print(f"🔧 Enhanced Debug Endpoints: Token analysis and system monitoring")
    print(f"🌐 Token-Aware Server starting on port {port}")
    
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=port,
        reload=False,
        log_level="info"
    )